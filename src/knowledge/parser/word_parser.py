"""
Word 文档解析器
支持解析 .docx 格式
"""

import os
from typing import Any, Dict, Optional

from .base import BaseParser, ParseResult, ParseStatus


class WordParser(BaseParser):
    """
    Word 文档解析器

    使用 python-docx 解析 .docx 文件
    """

    supported_extensions = ["doc", "docx"]
    name = "word"

    async def parse(self, file_path: str) -> ParseResult:
        """解析 Word 文档"""
        if not os.path.exists(file_path):
            return ParseResult(
                status=ParseStatus.FAILED,
                content="",
                error=f"File not found: {file_path}",
            )

        ext = file_path.lower().split(".")[-1]

        # .doc 格式需要特殊处理
        if ext == "doc":
            return await self._parse_doc(file_path)

        return await self._parse_docx(file_path)

    async def _parse_docx(self, file_path: str) -> ParseResult:
        """解析 .docx 文件"""
        try:
            from docx import Document
        except ImportError:
            return ParseResult(
                status=ParseStatus.FAILED,
                content="",
                error="python-docx not installed. Run: pip install python-docx",
            )

        try:
            doc = Document(file_path)

            content_parts = []
            metadata = {}
            sections = []
            tables = []

            # 提取核心属性
            core_props = doc.core_properties
            if core_props:
                metadata = {
                    "title": core_props.title,
                    "author": core_props.author,
                    "subject": core_props.subject,
                    "keywords": core_props.keywords,
                    "created": str(core_props.created) if core_props.created else None,
                    "modified": str(core_props.modified) if core_props.modified else None,
                }

            # 提取段落
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    content_parts.append(text)

                    # 检测标题
                    if para.style.name.startswith("Heading"):
                        sections.append({
                            "type": "heading",
                            "level": para.style.name,
                            "content": text,
                        })

            # 提取表格
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)

                tables.append({
                    "index": table_idx,
                    "rows": len(table_data),
                    "cols": len(table_data[0]) if table_data else 0,
                    "data": table_data[:5],  # 只保存前 5 行
                })

            content = "\n\n".join(content_parts)
            content = self._clean_content(content)

            return ParseResult(
                status=ParseStatus.SUCCESS,
                content=content,
                title=metadata.get("title"),
                author=metadata.get("author"),
                metadata=metadata,
                sections=sections,
                tables=tables,
            )

        except Exception as e:
            return ParseResult(
                status=ParseStatus.FAILED,
                content="",
                error=str(e),
            )

    async def _parse_doc(self, file_path: str) -> ParseResult:
        """解析 .doc 文件 (旧格式)"""
        # .doc 格式较难处理，建议使用 Tika 或转换工具
        return ParseResult(
            status=ParseStatus.PARTIAL,
            content="",
            error=".doc format is not directly supported. Please convert to .docx or use Tika.",
            warnings=["Legacy .doc format detected. Consider converting to .docx."],
        )
